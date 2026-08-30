"""
文档解析服务
支持 PDF、DOCX、TXT、Markdown 格式的文件解析
"""

from pathlib import Path

from app.config import SUPPORTED_EXTENSIONS


def parse_file(file_path: Path) -> str:
    """
    根据文件扩展名选择合适的解析器，提取文本内容

    Args:
        file_path: 待解析文件的路径

    Returns:
        str: 提取到的纯文本内容

    Raises:
        ValueError: 当文件格式不受支持时抛出
    """
    # 获取文件扩展名（小写）
    extension = file_path.suffix.lower()

    # 校验文件格式是否受支持
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件格式: {extension}，支持的格式: {SUPPORTED_EXTENSIONS}")

    if extension == ".pdf":
        return _parse_pdf(file_path)
    elif extension == ".docx":
        return _parse_docx(file_path)
    elif extension == ".txt":
        return _parse_txt(file_path)
    elif extension == ".md":
        return _parse_txt(file_path)  # Markdown 文件按纯文本读取

    raise ValueError(f"未知的文件格式: {extension}")


def _parse_pdf(file_path: Path) -> str:
    """
    解析 PDF 文件，提取文本内容

    Args:
        file_path: PDF 文件路径

    Returns:
        str: 提取到的文本内容
    """
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    text_parts = []

    # 遍历 PDF 的每一页，提取文本
    for page_num, page in enumerate(reader.pages, 1):
        page_text = page.extract_text()
        if page_text and page_text.strip():
            text_parts.append(page_text.strip())

    if not text_parts:
        raise ValueError("PDF 文件中未找到可提取的文本内容，可能是扫描版 PDF")

    return "\n\n".join(text_parts)


def _parse_docx(file_path: Path) -> str:
    """
    解析 DOCX 文件，提取文本内容
    同时遍历正文段落与表格（校园文档的费用标准、开放时间常以表格呈现，
    仅读 paragraphs 会全部丢失）

    Args:
        file_path: DOCX 文件路径

    Returns:
        str: 提取到的文本内容
    """
    from docx import Document

    doc = Document(str(file_path))
    text_parts = []

    # 遍历文档中的每个段落
    for paragraph in doc.paragraphs:
        if paragraph.text and paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # 遍历所有表格，逐行拼接单元格（以 | 分隔）
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                text_parts.append(row_text)

    if not text_parts:
        raise ValueError("DOCX 文件中未找到文本内容")

    return "\n\n".join(text_parts)


def _parse_txt(file_path: Path) -> str:
    """
    解析 TXT/MD 纯文本文件

    Args:
        file_path: 文本文件路径

    Returns:
        str: 文件内容
    """
    # 尝试多种编码读取（UTF-8 优先，GBK 作为中文 Windows 常见编码的备选）
    for encoding in ["utf-8", "gbk", "gb2312", "latin-1"]:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read()
            if content.strip():
                return content
        except (UnicodeDecodeError, UnicodeError):
            continue

    raise ValueError("无法读取文件，请检查文件编码（支持 UTF-8 / GBK）")
